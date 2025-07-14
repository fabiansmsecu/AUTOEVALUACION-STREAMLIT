# -----------------------------------
# PARTE 1
# -----------------------------------

# IMPORTACIONES
import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
from datetime import datetime
import plotly.express as px
import os
from openai import OpenAI

# -------------------------------
# CONFIGURAR DEEPSEEK
# -------------------------------
client = OpenAI(
    api_key=os.getenv("DEEPSEEK_API_KEY"),
    base_url="https://api.deepseek.com"
)

# -------------------------------
# RUBRICA COMPLETA (DESCRIPCIONES + RECOMENDACIONES)
# -------------------------------
# ✅ Aquí debes pegar el bloque de RÚBRICAS que generamos antes
# Para ahorrar espacio aquí, colócala desde el mensaje anterior
from rubricas import rubricas

# -------------------------------
# FUNCIONES
# -------------------------------

def procesar_calificaciones(calificaciones):
    """
    Calcula el promedio por dominio.
    """
    promedios = {
        aspecto: sum(calificacion for pregunta, calificacion in lista) / len(lista)
        for aspecto, lista in calificaciones.items()
    }
    return promedios

def generar_prompt_para_DeepSeek(calificaciones, promedios, empresa, rubro, tamanio):
    """
    Construye el prompt COMPLETO para DeepSeek,
    incluyendo preguntas, respuestas, descripciones y recomendaciones.
    """

    # Construir texto con respuestas completas
    texto_respuestas = ""

    for aspecto, preguntas in calificaciones.items():
        texto_respuestas += f"\n## {aspecto}\n"
        for pregunta, calificacion in preguntas:
            descripcion = rubricas[aspecto][pregunta][calificacion]['descripcion']
            recomendacion = rubricas[aspecto][pregunta][calificacion]['recomendacion']
            texto_respuestas += f"- **{pregunta}**\n    - Calificación: {calificacion}\n    - Descripción: {descripcion}\n    - Recomendación: {recomendacion}\n\n"

    # Armamos prompt extenso para IA
    prompt = f"""
Eres un auditor experto en la norma ISO 27001.

Estoy realizando una autoevaluación para la empresa "{empresa}", dedicada a "{rubro}", con aproximadamente {tamanio} empleados.

Estos son los resultados obtenidos:

{texto_respuestas}

Promedios por dominio:
{promedios}

Con base en esta información, genera un informe profesional que incluya:

- Hallazgos específicos sobre las áreas más críticas
- Identificación de los controles ISO 27001 que se incumplen o tienen debilidades
- Riesgos reales que estas debilidades implican para una empresa de este tamaño y sector
- Recomendaciones técnicas específicas y realistas
- Plan de acción priorizado, con tiempos estimados y responsables sugeridos
- Redacción profesional, clara y entendible para directivos

Evita recomendaciones genéricas. Redacta hallazgos y conclusiones personalizadas según los datos entregados.
"""

    return prompt

# -----------------------------------
# PARTE 2
# -----------------------------------

def generar_texto_DeepSeek(prompt):
    """
    Llama a DeepSeek para generar el texto profesional.
    """
    completion = client.chat.completions.create(
        model="deepseek-chat",
        messages=[
            {"role": "system", "content": "Eres un auditor experto en ISO 27001."},
            {"role": "user", "content": prompt}
        ]
    )

    texto_ia = completion.choices[0].message.content
    return texto_ia

def generar_informe_word(calificaciones, promedios, texto_ia,
                         nombre_compania, nombre_evaluador,
                         destinatario, fecha_evaluacion):
    """
    Genera un archivo Word con todos los resultados.
    """
    buffer = BytesIO()
    doc = Document()

    # Portada
    doc.add_heading('Informe de Evaluación de Cumplimiento de la Norma ISO 27001', 0)
    doc.add_paragraph(f'Compañía Evaluada: {nombre_compania}', style='Title')
    doc.add_paragraph(f'Evaluador: {nombre_evaluador}', style='Heading 3')
    doc.add_paragraph(f'Fecha de Evaluación: {fecha_evaluacion}', style='Heading 3')
    doc.add_page_break()

    # Dimensiones
    doc.add_heading('Dimensiones Evaluadas', level=1)
    dimensiones = {
        'Gestión de Acceso': "Evalúa la existencia y eficacia de políticas y procedimientos para la gestión de accesos.",
        'Seguridad Física y Ambiental': "Evalúa medidas físicas y ambientales para proteger los activos de TI.",
        'Gestión de Comunicaciones y Operaciones': "Evalúa prácticas para proteger datos en tránsito y operaciones seguras.",
        'Control de Acceso a la Información': "Evalúa controles para limitar acceso a información confidencial.",
        'Gestión de Incidentes de Seguridad de la Información': "Evalúa preparación y respuesta ante incidentes de seguridad."
    }
    for dim, desc in dimensiones.items():
        doc.add_heading(dim, level=2)
        doc.add_paragraph(desc)

    doc.add_page_break()

    # Resultados Detallados
    doc.add_heading('Resultados de la Evaluación', level=1)
    for aspecto, preguntas in calificaciones.items():
        doc.add_heading(aspecto, level=2)
        for pregunta, calificacion in preguntas:
            desc = rubricas[aspecto][pregunta][calificacion]['descripcion']
            recomendacion = rubricas[aspecto][pregunta][calificacion]['recomendacion']
            p = doc.add_paragraph()
            p.add_run(f"{pregunta}: ").bold = True
            p.add_run(f"{calificacion} - {desc}")
            doc.add_paragraph(f"Recomendación: {recomendacion}")
        doc.add_paragraph(f"Promedio del aspecto: {promedios[aspecto]:.2f} / 5")
        doc.add_paragraph()

    doc.add_page_break()

    # Texto IA
    doc.add_heading('Análisis y Recomendaciones Generadas por IA', level=1)
    doc.add_paragraph(texto_ia)

    # Guardar archivo en memoria
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generar_grafico(promedios):
    """
    Genera un gráfico de barras con los promedios.
    """
    df = pd.DataFrame(list(promedios.items()), columns=["Dominio", "Promedio"])
    fig = px.bar(
        df,
        x="Dominio",
        y="Promedio",
        color="Promedio",
        color_continuous_scale=["red", "yellow", "green"],
        range_y=[0, 5],
        title="Resultados por Dominio"
    )
    st.plotly_chart(fig)

# -------------------------------
# INTERFAZ STREAMLIT
# -------------------------------

def main():
    st.title("Evaluación ISO 27001 - Con IA DeepSeek")

    st.info("Completa todos los campos y responde el cuestionario para generar un informe detallado con análisis de IA.")

    nombre_compania = st.text_input("Nombre de la Compañía Evaluada", "")
    nombre_evaluador = st.text_input("Nombre del Evaluador", "")
    destinatario = st.text_input("Destinatario del Informe", "")
    rubro = st.text_input("Sector o Rubro de la Empresa (ej. salud, finanzas, retail, etc.)", "")
    tamanio = st.number_input("Cantidad aproximada de empleados", min_value=1, step=1)

    fecha_evaluacion = datetime.now().strftime("%d/%m/%Y %H:%M:%S")

    calificaciones = {key: [] for key in rubricas.keys()}

    for aspecto, preguntas in rubricas.items():
        st.subheader(aspecto)
        for pregunta, niveles in preguntas.items():
            opciones = [f"{k}: {v['descripcion']}" for k, v in niveles.items()]
            seleccion = st.selectbox(
                pregunta,
                opciones,
                key=f"{aspecto}_{pregunta}"
            )
            calificacion = int(seleccion.split(":")[0])
            calificaciones[aspecto].append((pregunta, calificacion))

    if st.button("Generar Informe"):
        if not all([nombre_compania, nombre_evaluador, destinatario, rubro, tamanio]):
            st.error("¡Completa todos los campos antes de generar el informe!")
        else:
            promedios = procesar_calificaciones(calificaciones)
            generar_grafico(promedios)

            # Generar prompt dinámico
            prompt = generar_prompt_para_DeepSeek(
                calificaciones,
                promedios,
                nombre_compania,
                rubro,
                tamanio
            )

            # Llamada a DeepSeek
            texto_ia = generar_texto_DeepSeek(prompt)

            # Generar informe Word
            buffer = generar_informe_word(
                calificaciones,
                promedios,
                texto_ia,
                nombre_compania,
                nombre_evaluador,
                destinatario,
                fecha_evaluacion
            )

            st.success("¡Informe generado exitosamente!")
            st.download_button(
                label="Descargar Informe en Word",
                data=buffer,
                file_name=f"Informe_ISO27001_{nombre_compania}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == "__main__":
    main()

