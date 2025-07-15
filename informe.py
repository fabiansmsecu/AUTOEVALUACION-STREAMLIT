from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from io import BytesIO
from rubricas import rubricas

def generar_conclusion(calificacion_final):
    if 0 <= calificacion_final <= 25:
        return "El departamento muestra un nivel muy bajo de cumplimiento. Se recomienda una revisión completa de los procesos de seguridad de la información."
    elif 26 <= calificacion_final <= 50:
        return "El departamento tiene un cumplimiento parcial. Existen controles implementados, pero no son suficientes ni consistentes."
    elif 51 <= calificacion_final <= 75:
        return "El departamento cumple en gran medida con la norma ISO 27001, aunque tiene áreas de mejora para alcanzar la excelencia."
    elif 76 <= calificacion_final <= 100:
        return "El departamento demuestra un excelente nivel de cumplimiento y aplica controles más allá de lo exigido por la norma."
    else:
        return "Calificación no válida."

def generar_informe_word(calificaciones, promedios_ponderados, calificacion_final,
                         nombre_compania, nombre_evaluador, destinatario, fecha_evaluacion):
    doc = Document()

    # Logo (opcional, si tienes logo.png en tu carpeta)
    try:
        doc.add_picture("logo.png", width=Inches(2))
    except:
        pass

    # Título grande y azul
    p = doc.add_paragraph()
    run = p.add_run('Informe de Evaluación de Cumplimiento ISO 27001')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph(f"Compañía Evaluada: {nombre_compania}")
    doc.add_paragraph(f"Evaluador: {nombre_evaluador}")
    doc.add_paragraph(f"Fecha de Evaluación: {fecha_evaluacion}")
    doc.add_paragraph()

    doc.add_heading('Resultados Detallados', level=1)

    for aspecto, preguntas in calificaciones.items():
        doc.add_heading(aspecto, level=2)

        # Crear tabla bonita
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Pregunta'
        hdr_cells[1].text = 'Calificación'
        hdr_cells[2].text = 'Descripción'
        hdr_cells[3].text = 'Recomendación'

        # Colorear cabecera de azul y texto blanco
        for cell in hdr_cells:
            cell._tc.get_or_add_tcPr().append(
                parse_xml(r'<w:shd {} w:fill="003366"/>'.format(nsdecls('w')))
            )
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True

        for pregunta, calificacion in preguntas:
            desc = rubricas[aspecto][pregunta][calificacion]['descripcion']
            recomendacion = rubricas[aspecto][pregunta][calificacion]['recomendacion']
            row_cells = table.add_row().cells
            row_cells[0].text = pregunta
            row_cells[1].text = str(calificacion)
            row_cells[2].text = desc
            row_cells[3].text = recomendacion

        doc.add_paragraph(f"Promedio del aspecto {aspecto}: {promedios_ponderados[aspecto]:.2f}/20")
        doc.add_paragraph()

    doc.add_heading('Calificación Final', level=1)
    doc.add_paragraph(f"Calificación total: {calificacion_final:.2f}/100")

    conclusion = generar_conclusion(calificacion_final)
    doc.add_heading('Conclusión', level=1)
    doc.add_paragraph(conclusion)

    # Guardar a BytesIO para poder devolverlo a Streamlit
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
